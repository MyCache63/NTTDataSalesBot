#!/usr/bin/env python3
"""Generate OneNTTData_SalesBot_Requirements_v03_April2026.docx
   Updated after April 2 demo with Shreeshan — incorporates confirmed details."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================================
# HELPERS
# ============================================================
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(24)

run = p.add_run('NTT DATA')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x1E, 0x50)

run = p.add_run('   |   ')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

run = p.add_run('kore')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x14, 0x64)
run = p.add_run('.ai')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0xBF, 0xA5)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('\n')
run.font.size = Pt(16)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p3.add_run('One NTT Data Sales Bot')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p4.add_run('Business Requirements Document')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x50, 0xD4)

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p5.add_run('v03 — 2026-04-02 14:00 PT')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p6.add_run('Agentic AI Presentation & Sales Enablement Platform\nfor NTT DATA Cloud Services Group')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p7.add_run('CONFIDENTIAL — For Internal Use Only')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph()
p8 = doc.add_paragraph()
p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p8.add_run('Prepared by: Kore.ai / NTT DATA Partnership Team\nApril 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# CHANGE LOG
# ============================================================
add_heading_styled('Change Log', 1)

cl_table = doc.add_table(rows=4, cols=3)
cl_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Version', 'Date', 'Changes']):
    cl_table.rows[0].cells[i].text = h
    for p in cl_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

cl_data = [
    ('v01', '2026-04-01', 'Initial requirements and design document'),
    ('v02', '2026-04-01', 'Separated requirements from design; added priority column; corrected stakeholder names'),
    ('v03', '2026-04-02', 'Updated after demo/discovery with Shreeshan (Apr 2). Confirmed: volume, implementation partner (Gradius), data sources (Confluence + Hot Spot portal), internal use, voice priority. Added: demo catalog concept, digital assistant use case, virtual teammate model. Moved answered questions to confirmed decisions.'),
]
for row_idx, (ver, date, changes) in enumerate(cl_data, 1):
    cl_table.rows[row_idx].cells[0].text = ver
    cl_table.rows[row_idx].cells[1].text = date
    cl_table.rows[row_idx].cells[2].text = changes

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Business Context',
    '3. Stakeholders',
    '4. Confirmed Decisions (from Discovery)',
    '5. Business Requirements',
    '6. Use Cases',
    '7. Data Requirements',
    '8. Remaining Open Questions',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading_styled('1. Executive Summary', 1)

doc.add_paragraph(
    'NTT DATA seeks an interactive, agentic AI system \u2014 the \u201cOne NTT Sales Bot\u201d \u2014 '
    'that serves as a virtual teammate for the NTT DATA sales organization. The system will '
    'join Microsoft Teams meetings as a participant, deliver pre-configured product presentations '
    'with voice narration, answer real-time Q&A from a continuously-updated knowledge base '
    '(Confluence, Hot Spot portal, product documentation), and act as an always-available digital '
    'assistant for product knowledge.'
)

doc.add_paragraph(
    'The bot will start with the Cloud Services Group and expand across the organization. '
    'Shreeshan Rangayan envisions a growing catalog of configurable demos \u2014 one per product or '
    'practice area \u2014 that any sales rep can launch on demand.'
)

doc.add_paragraph(
    'This document captures the business requirements gathered from discovery sessions on '
    'April 1\u20132, 2026 with Shreeshan Rangayan (NTT DATA) and the Kore.ai partnership team. '
    'It is intended as the foundation for subsequent design, scoping, and POC planning.'
)

# ============================================================
# 2. BUSINESS CONTEXT
# ============================================================
add_heading_styled('2. Business Context', 1)

add_heading_styled('The Problem', 2)
add_bullet('Product information is scattered across Confluence, Hot Spot portal, Seismic, slide decks, and brochures')
add_bullet('Sales reps spend significant time researching products before client calls')
add_bullet('Inconsistent messaging and positioning across the sales organization')
add_bullet('Subject-matter experts like Shreeshan are bottlenecks \u2014 reps come to them for answers that could be automated')
add_bullet('No interactive, on-demand way to get up to speed on products quickly')
add_bullet('New sales reps face a steep learning curve on the product portfolio')

add_heading_styled('The Vision', 2)
doc.add_paragraph(
    'A virtual teammate that knows everything about NTT DATA\u2019s cloud services portfolio. '
    'It can be added to Teams calls to present slides, answer real-time questions, and serve as '
    'a 24/7 digital assistant. When it doesn\u2019t know something, it logs the question so the knowledge '
    'base can be updated \u2014 creating a continuous improvement cycle.'
)

add_heading_styled('Business Value', 2)
add_bullet('Reduce research time for sales reps from hours to minutes', 'Time Savings: ')
add_bullet('Consistent, accurate product messaging across the entire sales org', 'Consistency: ')
add_bullet('New reps productive faster with AI-guided product knowledge', 'Onboarding: ')
add_bullet('Eliminates bottleneck on subject-matter experts for routine product questions', 'Scalability: ')
add_bullet('Available 24/7, across time zones, in multiple languages', 'Availability: ')
add_bullet('Continuously updated as Confluence and product documentation changes', 'Currency: ')

add_heading_styled('Target Users & Volume (Confirmed Apr 2)', 2)
add_bullet('200+ NTT DATA sales representatives in the Cloud Services Group')
add_bullet('Each user expected to run ~10 sessions per month (demos + Q&A)')
add_bullet('Estimated initial volume: 2,000+ sessions per month')
add_bullet('Usage includes both formal demo presentations and ad-hoc product Q&A')
add_bullet('Starting with One NTT Platform; expanding to FinOps, Cloud, and all other practices')

doc.add_page_break()

# ============================================================
# 3. STAKEHOLDERS
# ============================================================
add_heading_styled('3. Stakeholders', 1)

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Name', 'Role', 'Organization']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ('Shreeshan Rangayan', 'Project Lead', 'NTT DATA'),
    ('Wiley Oliver', 'Executive Sponsor', 'NTT DATA'),
    ('Gradius', 'Implementation Partner (Kore.ai vendor)', 'Kore.ai ecosystem'),
    ('Mike Ashe', 'Account Lead', 'Kore.ai'),
    ('Michael Dsousa', 'Relationship Manager', 'Kore.ai'),
    ('Deepak Anand', 'Solutions Architect', 'Kore.ai'),
]
for row_idx, (name, role, org) in enumerate(data, 1):
    table.rows[row_idx].cells[0].text = name
    table.rows[row_idx].cells[1].text = role
    table.rows[row_idx].cells[2].text = org

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Note: ')
run.bold = True
p.add_run(
    'Gradius is a Kore.ai vendor recently onboarded by NTT DATA. '
    'They currently support NTT DATA on MetLife and other existing Kore platform work. '
    'Many team members are former Kore.ai employees. Gradius will be responsible for building the solution.'
)

doc.add_page_break()

# ============================================================
# 4. CONFIRMED DECISIONS
# ============================================================
add_heading_styled('4. Confirmed Decisions (from Discovery)', 1)

doc.add_paragraph(
    'The following items were confirmed during the April 1\u20132 discovery sessions with Shreeshan Rangayan.'
)

decisions = [
    ('Usage Model', 'Internal use only. The bot is for NTT DATA\u2019s sales organization. Presentations may be delivered to external customers, but the tool itself is internal.'),
    ('Volume', '200+ users, each running ~10 sessions/month. Initial volume: 2,000+ sessions/month. Includes both formal demos and ad-hoc Q&A.'),
    ('Implementation Partner', 'Gradius (Kore.ai vendor) will build the solution. NTT DATA has already onboarded them as a vendor.'),
    ('Voice', 'Voice narration is critical and required. Approach: start with text-based interaction, then layer in voice. Voice responses should be concise \u2014 for lengthy answers, the bot should offer to send a document or email rather than speaking at length.'),
    ('Data Sources', 'Three primary sources: (1) Confluence \u2014 all technical documentation, (2) Hot Spot \u2014 NTT DATA\u2019s internal pre-sales support portal, (3) PowerPoint decks with associated scripts and documents. CRM is Salesforce.'),
    ('Demo Catalog', 'The system should support a growing catalog of configurable demos \u2014 one per product/practice (e.g., One NTT Platform, NTT Composable, NTT Flash, Cloud, FinOps). Any sales rep can browse the catalog and launch a demo.'),
    ('Virtual Teammate Model', 'The bot is envisioned as a virtual meeting participant added to Teams calls. It listens to the conversation, can be addressed by name, presents slides on command, and answers Q&A in real time.'),
    ('Unanswered Questions', 'When the bot cannot answer a question, it acknowledges the gap, logs the question, and a human updates Confluence. The bot picks up the new content on its next sync \u2014 creating a continuous improvement loop.'),
    ('Starting Scope', 'Begin with the One NTT Platform product line. Expand to other practices once proven.'),
    ('Sample Data', 'Shreeshan will provide sample PowerPoints and associated documents to support the next phase of the POC.'),
]

dec_table = doc.add_table(rows=len(decisions)+1, cols=2)
dec_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Topic', 'Decision']):
    dec_table.rows[0].cells[i].text = h
    for p in dec_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (topic, decision) in enumerate(decisions, 1):
    dec_table.rows[row_idx].cells[0].text = topic
    dec_table.rows[row_idx].cells[1].text = decision

doc.add_page_break()

# ============================================================
# 5. BUSINESS REQUIREMENTS
# ============================================================
add_heading_styled('5. Business Requirements', 1)

doc.add_paragraph(
    'The following requirements were provided by Shreeshan Rangayan across the April 1\u20132 '
    'discovery sessions. Requirements updated in v03 are marked with \u2605.'
)

reqs = [
    ('REQ-01', 'Agentic AI Presenter',
     'Create subject-specific Agentic AI presenters (e.g., \u201cOne NTT Platform Specialist\u201d) that can deliver presentations interactively.',
     'Must'),
    ('REQ-02', 'Supporting Material Integration',
     'Ingest and utilize One NTT slides, presentation scripts, and Confluence content as the knowledge base.',
     'Must'),
    ('REQ-03', 'Demo Catalog \u2605',
     'Maintain a browsable catalog of configurable demos organized by product/practice area (e.g., One NTT Platform, NTT Composable, FinOps, Cloud). Any sales rep can select and launch a demo from the catalog. The catalog grows over time as new demos are added.',
     'Must'),
    ('REQ-04', 'Microsoft Teams Integration',
     'Add the Agentic AI agent into Microsoft Teams as a meeting participant for live presentation delivery and real-time Q&A.',
     'Must'),
    ('REQ-05', 'Live Presentation Delivery',
     'Start the Agentic AI to give a presentation to an audience in real-time, narrating slides with a configurable voice.',
     'Must'),
    ('REQ-06', 'Interactive Q&A',
     'Anyone in the audience can interrupt with questions during the presentation. The agent acknowledges and answers from the script, PPT content, or Confluence.',
     'Must'),
    ('REQ-07', 'Graceful Handling of Unknowns \u2605',
     'When the agent cannot answer a question, it acknowledges the gap, says it will follow up, and captures the question in meeting minutes. These logged questions feed a continuous improvement cycle \u2014 humans update Confluence, and the bot picks up new content on its next sync.',
     'Must'),
    ('REQ-08', 'Confluence Content Sync \u2605',
     'Scheduled synchronization with Confluence (e.g., daily). Detect and handle new, deleted, and modified content automatically so the bot is always current with technical documentation.',
     'Must'),
    ('REQ-09', 'Presentation Time Estimation',
     'Ability to estimate presentation duration based on the number and content of standard slides.',
     'Should'),
    ('REQ-10', 'Translation and Accent Selection',
     'Ability to translate presentation content and choose voice accent for multi-language delivery.',
     'Should'),
    ('REQ-11', 'Meeting Minutes Generation',
     'Automatically generate and send meeting minutes after a presentation session, including any unanswered questions flagged for follow-up.',
     'Must'),
    ('REQ-12', 'Presentation Version Control',
     'Maintain version history and change tracking for presentations in the catalog.',
     'Should'),
    ('REQ-13', 'Multi-Source Content Connection \u2605',
     'Connect content from: (1) PowerPoint slides and notes sections, (2) associated Word documents and scripts, (3) Confluence technical documentation, (4) Hot Spot pre-sales support portal.',
     'Must'),
    ('REQ-14', 'Usage Reporting',
     'Track and report the number of presentations conducted, by topic, presenter, and audience.',
     'Should'),
    ('REQ-15', 'CSAT Surveys',
     'Send out CSAT surveys after presentations and collect audience feedback.',
     'Should'),
    ('REQ-16', 'Sentiment Analysis',
     'Capture sentiment scores based on conversation dynamics during presentations.',
     'Should'),
    ('REQ-17', 'Voice Input \u2605 (New)',
     'Users can ask questions and give commands via voice instead of typing. The bot should support speech-to-text input in addition to text chat.',
     'Must'),
    ('REQ-18', 'Concise Voice Responses \u2605 (New)',
     'When responding by voice, the bot should give brief, pointed answers. For detailed or lengthy responses, it should offer to send a document or email rather than speaking at length.',
     'Must'),
    ('REQ-19', 'Digital Assistant Mode \u2605 (New)',
     'Beyond formal demos, the bot should serve as an always-available digital assistant that any sales rep can query for quick product questions without launching a full presentation.',
     'Must'),
    ('REQ-20', 'Hot Spot Portal Integration \u2605 (New)',
     'Integrate with Hot Spot (NTT DATA\u2019s internal pre-sales support portal) so the bot can be accessed as a default capability within the existing pre-sales workflow.',
     'Should'),
]

table = doc.add_table(rows=len(reqs)+1, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['ID', 'Requirement', 'Description', 'Priority']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (rid, name, desc, priority) in enumerate(reqs, 1):
    table.rows[row_idx].cells[0].text = rid
    table.rows[row_idx].cells[1].text = name
    table.rows[row_idx].cells[2].text = desc
    table.rows[row_idx].cells[3].text = priority

doc.add_page_break()

# ============================================================
# 6. USE CASES
# ============================================================
add_heading_styled('6. Use Cases', 1)

use_cases = [
    ('UC-01', 'Sales Rep Self-Enablement',
     'A sales rep preparing for a customer meeting opens the Sales Bot and asks: '
     '\u201cI need to sell cloud migration services to a financial services company. '
     'What should I position and which slides should I use?\u201d '
     'The bot responds with a tailored positioning strategy and recommends specific slides from the corpus.',
     ['REQ-01', 'REQ-02', 'REQ-06', 'REQ-13', 'REQ-19']),
    ('UC-02', 'Live AI-Assisted Presentation (Virtual Teammate)',
     'In a Microsoft Teams meeting with a customer, the AI agent is added as a participant. '
     'A sales leader says \u201cPresent the One NTT Cloud Platform overview.\u201d The agent narrates '
     'the presentation with voice. A customer interrupts: \u201cWhat LLM do you use?\u201d The agent '
     'answers from Confluence. If it can\u2019t answer, it says \u201cGreat question \u2014 let me note that '
     'and we\u2019ll follow up\u201d and logs the question.',
     ['REQ-04', 'REQ-05', 'REQ-06', 'REQ-07', 'REQ-09', 'REQ-17', 'REQ-18']),
    ('UC-03', 'Digital Assistant Q&A',
     'A sales rep needs a quick answer before a call. They open the bot and ask: '
     '\u201cDoes NTT DATA have FedRAMP authorization?\u201d The bot gives a concise answer with sources. '
     'No presentation needed \u2014 just fast, accurate product knowledge on demand. '
     'This replaces the current pattern of messaging Shreeshan or other SMEs for quick answers.',
     ['REQ-01', 'REQ-08', 'REQ-13', 'REQ-17', 'REQ-19']),
    ('UC-04', 'Demo Catalog Browse & Launch',
     'A sales engineer opens the demo catalog and sees available demos: One NTT Platform, '
     'NTT Composable, NTT Flash, Cloud Services, FinOps. They select \u201cFinOps\u201d and configure '
     'a 10-minute version. The bot generates the presentation and is ready to deliver it in '
     'the next Teams call.',
     ['REQ-03', 'REQ-05', 'REQ-09', 'REQ-12']),
    ('UC-05', 'New Hire Onboarding',
     'A new sales hire uses the bot to rapidly learn the product portfolio. They ask open-ended '
     'questions, request presentations on specific topics, and get coaching on how to position products '
     'for different industries and competitive scenarios.',
     ['REQ-01', 'REQ-03', 'REQ-06', 'REQ-13', 'REQ-19']),
    ('UC-06', 'Post-Presentation Follow-Up',
     'After a live presentation, the system automatically generates meeting minutes capturing '
     'what was presented, questions asked, answers given, and any unanswered items flagged for '
     'follow-up. A CSAT survey is sent to attendees.',
     ['REQ-07', 'REQ-11', 'REQ-15', 'REQ-16']),
    ('UC-07', 'Knowledge Gap Closure (Continuous Improvement)',
     'After a demo, the bot logged 3 unanswered questions. A product manager reviews the log, '
     'updates Confluence with the answers. The next day, the bot\u2019s scheduled Confluence sync '
     'picks up the new content. The next time someone asks those questions, the bot answers correctly.',
     ['REQ-07', 'REQ-08']),
]

for uc_id, title, desc, req_refs in use_cases:
    add_heading_styled(f'{uc_id}: {title}', 2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    run = p.add_run('Related Requirements: ')
    run.bold = True
    p.add_run(', '.join(req_refs))

doc.add_page_break()

# ============================================================
# 7. DATA REQUIREMENTS
# ============================================================
add_heading_styled('7. Data Requirements', 1)

doc.add_paragraph(
    'The following data is needed from NTT DATA to support the system:'
)

add_heading_styled('Immediate Need (for POC)', 2)
doc.add_paragraph(
    'Shreeshan confirmed he will send sample materials. We need:'
)
add_bullet('20\u201330 PowerPoint decks covering key Cloud Services Group products')
add_bullet('Associated documents for those decks (scripts, narratives, Word docs)')
add_bullet('Product brochures and one-pagers')

add_heading_styled('Production Data Sources (Confirmed Apr 2)', 2)

add_bullet('All technical documentation for Cloud Services Group', 'Confluence: ')
add_bullet('Pre-sales support portal \u2014 eventual integration target for the bot', 'Hot Spot Portal: ')
add_bullet('CRM \u2014 potential source for customer context in future phases', 'Salesforce: ')
add_bullet('Content library (if still in use alongside Hot Spot)', 'Seismic: ')
add_bullet('All product slide decks with notes sections and associated scripts', 'Slide Decks: ')

add_heading_styled('Additional Production Data (as scope expands)', 2)
add_bullet('Competitive positioning documents')
add_bullet('Case studies and customer success stories')
add_bullet('Technical architecture documents')

doc.add_page_break()

# ============================================================
# 8. REMAINING OPEN QUESTIONS
# ============================================================
add_heading_styled('8. Remaining Open Questions', 1)

doc.add_paragraph(
    'The following questions remain open after the April 2 discovery session. '
    'Items resolved in this session have been moved to Section 4 (Confirmed Decisions).'
)

questions = [
    ('Pricing Model', 'Per-session cost, platform license, or managed service engagement? How does this align with the existing $3.5M budget framework? (Mike noted this will be addressed later.)'),
    ('Security & Compliance', 'What are the data handling, security, and compliance requirements for product documentation in this system? Any restrictions on LLM providers or data residency?'),
    ('Confluence Scope', 'Which Confluence spaces are in scope for the initial build? Who grants access? How frequently does content change?'),
    ('Teams Deployment Model', 'What is the Teams deployment model? Is there an existing bot framework or app catalog in use? Any IT approval process for adding bots to Teams?'),
    ('Success Metrics', 'How will NTT DATA measure success? Time saved per rep? Presentation quality? User adoption rates? Reduction in SME interruptions?'),
    ('Competitive Landscape', 'Are other vendors being evaluated for this solution? What evaluation criteria are being used?'),
    ('Hot Spot Integration Details', 'What technology stack does Hot Spot run on? What level of integration is expected \u2014 embedded iframe, API, or full native integration?'),
    ('Sample Data Timeline', 'When will Shreeshan send the sample PowerPoints and associated documents? (He confirmed he would \u2014 need to follow up on timing.)'),
]

table = doc.add_table(rows=len(questions)+1, cols=2)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Topic', 'Question / Notes']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (topic, question) in enumerate(questions, 1):
    table.rows[row_idx].cells[0].text = topic
    table.rows[row_idx].cells[1].text = question

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/michaelashe/Documents/KoreClients/NTT_Data_SalesBot/OneNTTData_SalesBot_Requirements_v03_April2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
