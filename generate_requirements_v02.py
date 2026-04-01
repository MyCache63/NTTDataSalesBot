#!/usr/bin/env python3
"""Generate OneNTTData_SalesBot_Requirements_v02_April2026.docx — Requirements only, no design speculation."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================================
# HELPERS
# ============================================================
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(24)

run = p.add_run('NTT DATA')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x1E, 0x50)

run = p.add_run('   |   ')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

run = p.add_run('kore')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x14, 0x64)
run = p.add_run('.ai')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0xBF, 0xA5)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('\n')
run.font.size = Pt(16)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p3.add_run('One NTT Data Sales Bot')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p4.add_run('Business Requirements Document')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x50, 0xD4)

doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p5.add_run('v02 — 2026-04-01 17:30 PT')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p6.add_run('Agentic AI Presentation Platform\nfor NTT DATA Cloud Services Group')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p7.add_run('CONFIDENTIAL — For Internal Use Only')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_paragraph()
p8 = doc.add_paragraph()
p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p8.add_run('Prepared by: Kore.ai / NTT DATA Partnership Team\nApril 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Business Context',
    '3. Stakeholders',
    '4. Business Requirements',
    '5. Use Cases',
    '6. Data Requirements',
    '7. Open Questions for Discovery',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading_styled('1. Executive Summary', 1)

doc.add_paragraph(
    'NTT DATA seeks an interactive, agentic AI system — the "One NTT Sales Bot" — '
    'that enables any member of the NTT DATA sales organization to access on-demand, intelligent '
    'presentations about NTT DATA\'s cloud services portfolio. The system will serve as a virtual '
    'subject-matter expert capable of presenting slides, answering real-time Q&A, and drawing from '
    'a continuously-updated knowledge base (Confluence, Seismic, product documentation).'
)

doc.add_paragraph(
    'This document captures the business requirements gathered from discovery sessions with '
    'Shreeshan Rangayan (NTT DATA) and the Kore.ai partnership team. It is intended as the '
    'foundation for subsequent design and scoping discussions.'
)

# ============================================================
# 2. BUSINESS CONTEXT
# ============================================================
add_heading_styled('2. Business Context', 1)

add_heading_styled('The Problem', 2)
add_bullet('Product information is scattered across Seismic, Confluence, slide decks, and brochures')
add_bullet('Sales reps spend significant time researching products before client calls')
add_bullet('Inconsistent messaging and positioning across the sales organization')
add_bullet('No interactive, on-demand way to get up to speed on products quickly')
add_bullet('New sales reps face a steep learning curve on the product portfolio')

add_heading_styled('The Vision', 2)
doc.add_paragraph(
    'An always-available AI specialist that knows everything about NTT DATA\'s cloud services portfolio '
    'and can present, discuss, and answer questions interactively — both for internal sales enablement '
    'and potentially for customer-facing demonstrations.'
)

add_heading_styled('Business Value', 2)
add_bullet('Reduce research time for sales reps from hours to minutes', 'Time Savings: ')
add_bullet('Consistent, accurate product messaging across the entire sales org', 'Consistency: ')
add_bullet('New reps productive faster with AI-guided product knowledge', 'Onboarding: ')
add_bullet('Available 24/7, across time zones, in multiple languages', 'Availability: ')
add_bullet('Continuously updated as product documentation changes', 'Currency: ')

add_heading_styled('Target Users', 2)
add_bullet('NTT DATA sales representatives (estimated 200-500 people)')
add_bullet('Sales engineers preparing for customer engagements')
add_bullet('New hires onboarding to the product portfolio')
add_bullet('Potentially: customers receiving AI-assisted product demonstrations')

doc.add_page_break()

# ============================================================
# 3. STAKEHOLDERS
# ============================================================
add_heading_styled('3. Stakeholders', 1)

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Name', 'Role', 'Organization']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ('Shreeshan Rangayan', 'Project Lead', 'NTT DATA'),
    ('Wiley Oliver', 'Executive Sponsor', 'NTT DATA'),
    ('Mike Ashe', 'Account Lead', 'Kore.ai'),
    ('Michael Dsousa', 'Relationship Manager', 'Kore.ai'),
    ('Deepak Anand', 'Solutions Architect', 'Kore.ai'),
]
for row_idx, (name, role, org) in enumerate(data, 1):
    table.rows[row_idx].cells[0].text = name
    table.rows[row_idx].cells[1].text = role
    table.rows[row_idx].cells[2].text = org

doc.add_paragraph()
doc.add_page_break()

# ============================================================
# 4. BUSINESS REQUIREMENTS
# ============================================================
add_heading_styled('4. Business Requirements', 1)

doc.add_paragraph(
    'The following requirements were provided by Shreeshan Rangayan and represent the full scope '
    'of the envisioned system.'
)

reqs = [
    ('REQ-01', 'Agentic AI Presenter',
     'Create subject-specific Agentic AI presenters (e.g., "One NTT Platform Specialist") that can deliver presentations interactively.',
     'Must'),
    ('REQ-02', 'Supporting Material Integration',
     'Ingest and utilize One NTT slides, presentation scripts, and Confluence content as the knowledge base.',
     'Must'),
    ('REQ-03', 'Catalogue Access',
     'Anyone in NTT should be able to open a presentation from a catalogue of available topics.',
     'Must'),
    ('REQ-04', 'Microsoft Teams Integration',
     'Add the Agentic AI agent into Microsoft Teams as a member for live presentation delivery.',
     'Must'),
    ('REQ-05', 'Live Presentation Delivery',
     'Start the Agentic AI to give a presentation to an audience in real-time.',
     'Must'),
    ('REQ-06', 'Interactive Q&A',
     'Anyone in the audience can ask questions during the presentation. The agent acknowledges and answers from the script, PPT content, or Confluence.',
     'Must'),
    ('REQ-07', 'Graceful Handling of Unknowns',
     'When the agent cannot answer a question, it apologizes, acknowledges the gap, and captures the question in meeting minutes for follow-up.',
     'Must'),
    ('REQ-08', 'Confluence Content Sync',
     'Ability to refresh updates from Confluence — detect and handle new, deleted, and modified content automatically.',
     'Must'),
    ('REQ-09', 'Presentation Time Estimation',
     'Ability to estimate presentation duration based on the number and content of standard slides.',
     'Should'),
    ('REQ-10', 'Translation and Accent Selection',
     'Ability to translate presentation content and choose voice accent for multi-language delivery.',
     'Should'),
    ('REQ-11', 'Meeting Minutes Generation',
     'Automatically generate and send meeting minutes after a presentation session, including any unanswered questions.',
     'Must'),
    ('REQ-12', 'Presentation Catalogue with Version Control',
     'Maintain a catalogue of available presentations with version history and change tracking.',
     'Should'),
    ('REQ-13', 'Multi-Source Content Connection',
     'Connect content from PowerPoint notes sections, attached Word documents, and internal portals.',
     'Must'),
    ('REQ-14', 'Usage Reporting',
     'Track and report the number of presentations conducted, by topic, presenter, and audience.',
     'Should'),
    ('REQ-15', 'CSAT Surveys',
     'Send out CSAT surveys after presentations and collect audience feedback.',
     'Should'),
    ('REQ-16', 'Sentiment Analysis',
     'Capture sentiment scores based on conversation dynamics during presentations.',
     'Should'),
]

table = doc.add_table(rows=len(reqs)+1, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['ID', 'Requirement', 'Description', 'Priority']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (rid, name, desc, priority) in enumerate(reqs, 1):
    table.rows[row_idx].cells[0].text = rid
    table.rows[row_idx].cells[1].text = name
    table.rows[row_idx].cells[2].text = desc
    table.rows[row_idx].cells[3].text = priority

doc.add_page_break()

# ============================================================
# 5. USE CASES
# ============================================================
add_heading_styled('5. Use Cases', 1)

use_cases = [
    ('UC-01', 'Sales Rep Self-Enablement',
     'A sales rep preparing for a customer meeting opens the Sales Bot and asks: '
     '"I need to sell cloud migration services to a financial services company. '
     'What should I position and which slides should I use?" '
     'The bot responds with a tailored positioning strategy and recommends specific slides from the corpus.',
     ['REQ-01', 'REQ-02', 'REQ-06', 'REQ-13']),
    ('UC-02', 'Live AI-Assisted Presentation',
     'In a Microsoft Teams meeting, the AI agent is added as a participant. A sales leader says '
     '"Present the One NTT Cloud Platform overview." The agent narrates the presentation with voice. '
     'Audience members interrupt with questions and the agent answers in real-time.',
     ['REQ-04', 'REQ-05', 'REQ-06', 'REQ-07', 'REQ-09']),
    ('UC-03', 'Customer-Facing Product Demo',
     'A potential customer interacts with the Sales Bot through a web interface or Teams. '
     'They ask questions about NTT DATA\'s cloud services and capabilities. '
     'The bot provides accurate, consistent answers drawn from the latest product documentation.',
     ['REQ-01', 'REQ-02', 'REQ-06', 'REQ-08']),
    ('UC-04', 'New Hire Onboarding',
     'A new sales hire uses the bot to rapidly learn the product portfolio. They ask open-ended '
     'questions, request presentations on specific topics, and get coaching on how to position products '
     'for different industries and competitive scenarios.',
     ['REQ-01', 'REQ-03', 'REQ-06', 'REQ-13']),
    ('UC-05', 'Post-Presentation Follow-Up',
     'After a live presentation, the system automatically generates meeting minutes capturing '
     'what was presented, questions asked, answers given, and any unanswered items. '
     'A CSAT survey is sent to attendees.',
     ['REQ-07', 'REQ-11', 'REQ-15', 'REQ-16']),
]

for uc_id, title, desc, req_refs in use_cases:
    add_heading_styled(f'{uc_id}: {title}', 2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    run = p.add_run('Related Requirements: ')
    run.bold = True
    p.add_run(', '.join(req_refs))

doc.add_page_break()

# ============================================================
# 6. DATA REQUIREMENTS
# ============================================================
add_heading_styled('6. Data Requirements', 1)

doc.add_paragraph(
    'The following data is needed from NTT DATA to support the system:'
)

add_heading_styled('Minimum for Initial Build', 2)
add_bullet('20-30 PowerPoint decks covering key Cloud Services Group products')
add_bullet('Product brochures and one-pagers')
add_bullet('Existing presentation scripts or speaker notes')
add_bullet('Read-only access to relevant Confluence spaces')

add_heading_styled('Full Production Data', 2)
add_bullet('Complete Confluence space access for Cloud Services documentation')
add_bullet('Seismic content library access or export')
add_bullet('All product slide decks with notes sections')
add_bullet('Competitive positioning documents')
add_bullet('Case studies and customer success stories')
add_bullet('Technical architecture documents')

doc.add_page_break()

# ============================================================
# 7. OPEN QUESTIONS FOR DISCOVERY
# ============================================================
add_heading_styled('7. Open Questions for Discovery', 1)

doc.add_paragraph(
    'The following questions need to be resolved through further discovery with Shreeshan Rangayan '
    'and the NTT DATA team.'
)

questions = [
    ('Volume & Usage', 'How many users and sessions per month are expected? Current estimate: 200-500 users in the sales org, frequency TBD.'),
    ('Pricing Model', 'Is this a per-session cost, platform license, or managed service engagement? How does this align with the existing $3.5M budget framework?'),
    ('Build vs. Deliver', 'Does NTT DATA want to build on the Kore.ai platform themselves, or have Kore.ai build and deliver a turnkey solution?'),
    ('Data Availability', 'When can NTT DATA provide the initial 20-30 sample documents for the first build? Who is the data owner?'),
    ('Internal vs. External', 'Is this strictly for internal sales enablement, or also intended for customer-facing use? Discovery suggests both.'),
    ('Voice Priority', 'How critical is voice narration for the initial delivery? Is a text-based interactive agent sufficient to start?'),
    ('Competitive Landscape', 'Are other vendors being evaluated for this solution? What evaluation criteria are being used?'),
    ('Success Metrics', 'How will NTT DATA measure success? Time saved per rep? Presentation quality? User adoption rates?'),
    ('Security & Compliance', 'What are the data handling, security, and compliance requirements for product documentation in this system?'),
    ('Confluence Scope', 'Which Confluence spaces are in scope? Who grants access? How frequently does content change?'),
    ('Teams Deployment', 'What is the Teams deployment model? Is there an existing bot framework or app catalog in use?'),
]

table = doc.add_table(rows=len(questions)+1, cols=2)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Topic', 'Question / Notes']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (topic, question) in enumerate(questions, 1):
    table.rows[row_idx].cells[0].text = topic
    table.rows[row_idx].cells[1].text = question

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/michaelashe/Documents/NTT_Data_SalesBot/OneNTTData_SalesBot_Requirements_v02_April2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
