#!/usr/bin/env python3
"""Generate OneNTTData_SalesBot_RequirementsAndDesign_v01_April2026.docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================================
# TITLE PAGE
# ============================================================
# Logos line (text-based since we can't easily embed SVG)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(24)

# NTT DATA logo text
run = p.add_run('NTT DATA')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x1E, 0x50)

run = p.add_run('   |   ')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# Kore.ai logo text
run = p.add_run('kore')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x14, 0x64)
run = p.add_run('.ai')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0xBF, 0xA5)

# Spacing
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p2.add_run('\n')
run.font.size = Pt(16)

# Title
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p3.add_run('One NTT Data Sales Bot')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p4.add_run('Requirements & Design Document')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x50, 0xD4)

# Spacing
doc.add_paragraph()

# Version stamp
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p5.add_run('v01 — 2026-04-01 15:30 PT')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# Subtitle
p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p6.add_run('Agentic AI Presentation Platform\nfor NTT DATA Cloud Services Group')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

# Confidentiality
p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p7.add_run('CONFIDENTIAL — For Internal Use Only')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

# Prepared by
doc.add_paragraph()
p8 = doc.add_paragraph()
p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p8.add_run('Prepared by: Kore.ai / NTT DATA Partnership Team\nApril 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Page break
doc.add_page_break()

# ============================================================
# Helper functions
# ============================================================
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading_styled('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Business Context & Opportunity',
    '3. Business Requirements (from Shreeshan / NTT DATA)',
    '4. Use Cases',
    '5. Solution Architecture',
    '6. Delivery Approach — RAG vs. Embedded Prompt',
    '7. Technical Specifications',
    '8. Data Requirements',
    '9. Integration Points',
    '10. MVP / POC Scope',
    '11. Future Roadmap',
    '12. Open Questions',
    '13. Appendix — NTT DATA Cloud Services Knowledge Base',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading_styled('1. Executive Summary', 1)

doc.add_paragraph(
    'NTT DATA seeks an interactive, agentic AI presentation system — the "One NTT Sales Bot" — '
    'that enables any member of the NTT DATA sales organization to access on-demand, intelligent '
    'presentations about NTT DATA\'s cloud services portfolio. The system will serve as a virtual '
    'subject-matter expert, capable of presenting slides, answering real-time Q&A, and drawing from '
    'a continuously-updated knowledge base (Confluence, Seismic, product documentation).'
)

doc.add_paragraph(
    'This document captures the requirements gathered from discovery sessions with Shreeshan (NTT DATA) '
    'and the Kore.ai partnership team, defines the solution architecture, and outlines the POC/MVP approach.'
)

add_heading_styled('Key Stakeholders', 2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['Name', 'Role', 'Organization']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ('Shreeshan', 'Sponsor / Decision Maker', 'NTT DATA'),
    ('Wiley', 'Decision Maker (Cloud Services Org)', 'NTT DATA'),
    ('Mike Ashe', 'Account Lead', 'Kore.ai'),
    ('Michael Dsousa', 'Relationship Manager', 'Kore.ai'),
    ('Deepak Anand', 'Solutions Architect', 'Kore.ai'),
]
for row_idx, (name, role, org) in enumerate(data, 1):
    table.rows[row_idx].cells[0].text = name
    table.rows[row_idx].cells[1].text = role
    table.rows[row_idx].cells[2].text = org

doc.add_paragraph()

# ============================================================
# 2. BUSINESS CONTEXT
# ============================================================
add_heading_styled('2. Business Context & Opportunity', 1)

doc.add_paragraph(
    'NTT DATA\'s sales organization currently faces fragmented access to product information. '
    'Sales reps must manually research products across Seismic, Confluence, slide decks, and '
    'brochures before client engagements. This is time-consuming and results in inconsistent '
    'positioning across the organization.'
)

add_heading_styled('The Problem', 2)
add_bullet('Product information is scattered across Seismic, Confluence, slide decks, and brochures')
add_bullet('Sales reps spend significant time researching products before client calls')
add_bullet('Inconsistent messaging and positioning across the sales organization')
add_bullet('No interactive, on-demand way to get up to speed on products quickly')
add_bullet('New sales reps face a steep learning curve on the product portfolio')

add_heading_styled('The Vision', 2)
doc.add_paragraph(
    'An always-available AI specialist that knows everything about NTT DATA\'s cloud services portfolio '
    'and can present, discuss, and answer questions interactively — both for internal sales enablement '
    'and potentially for customer-facing demonstrations.'
)

add_heading_styled('Business Value', 2)
add_bullet('Reduce research time for sales reps from hours to minutes', 'Time Savings: ')
add_bullet('Consistent, accurate product messaging across the entire sales org', 'Consistency: ')
add_bullet('New reps productive faster with AI-guided product knowledge', 'Onboarding: ')
add_bullet('Available 24/7, across time zones, in multiple languages', 'Availability: ')
add_bullet('Continuously updated as product documentation changes', 'Currency: ')

add_heading_styled('Target Users', 2)
add_bullet('NTT DATA sales representatives (200-500 people in the sales organization)')
add_bullet('Sales engineers preparing for customer engagements')
add_bullet('New hires onboarding to the product portfolio')
add_bullet('Potentially: customers receiving AI-assisted product demonstrations')

doc.add_page_break()

# ============================================================
# 3. BUSINESS REQUIREMENTS
# ============================================================
add_heading_styled('3. Business Requirements (from Shreeshan / NTT DATA)', 1)

doc.add_paragraph(
    'The following requirements were provided by Shreeshan and represent the full scope of the envisioned system. '
    'The POC will address a subset (see Section 10).'
)

reqs = [
    ('REQ-01', 'Agentic AI Presenter', 'Create a subject-specific Agentic AI presenter (e.g., "One NTT Platform Specialist") that can deliver presentations interactively.'),
    ('REQ-02', 'Supporting Material Integration', 'Ingest and utilize One NTT slides, presentation scripts, and Confluence content as the knowledge base.'),
    ('REQ-03', 'Catalogue Access', 'Anyone in NTT should be able to open a presentation from a catalogue of available topics.'),
    ('REQ-04', 'Teams Integration', 'Add the Kore Agentic AI agent into Microsoft Teams as a member for live presentation delivery.'),
    ('REQ-05', 'Live Presentation', 'Start the Agentic AI to give a presentation to an audience in real-time.'),
    ('REQ-06', 'Interactive Q&A', 'Anyone in the audience can ask questions during the presentation; the agent acknowledges and answers from the script, PPT, or Confluence.'),
    ('REQ-07', 'Graceful Unknowns', 'When the agent cannot answer a question, it apologizes and captures the question in meeting minutes for follow-up.'),
    ('REQ-08', 'Confluence Sync', 'Ability to refresh updates from Confluence — handle new, deleted, and modified content.'),
    ('REQ-09', 'Time Estimation', 'Ability to estimate presentation time based on standard slides.'),
    ('REQ-10', 'Translation & Accent', 'Ability to translate content and choose voice accent for multi-language delivery.'),
    ('REQ-11', 'Meeting Minutes', 'Automatically generate and send meeting minutes after a presentation session.'),
    ('REQ-12', 'Presentation Catalogue', 'Ability to create and maintain a catalogue of presentations with version control.'),
    ('REQ-13', 'Multi-Source Content', 'Connect content from notes sections, attached Word docs, and internal portals.'),
    ('REQ-14', 'Usage Reporting', 'Track and report the number of presentations conducted.'),
    ('REQ-15', 'CSAT Surveys', 'Send out CSAT surveys and collect audience feedback after presentations.'),
    ('REQ-16', 'Sentiment Analysis', 'Capture sentiment scores based on conversation during presentations.'),
]

table = doc.add_table(rows=len(reqs)+1, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['ID', 'Requirement', 'Description']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (rid, name, desc) in enumerate(reqs, 1):
    table.rows[row_idx].cells[0].text = rid
    table.rows[row_idx].cells[1].text = name
    table.rows[row_idx].cells[2].text = desc

doc.add_page_break()

# ============================================================
# 4. USE CASES
# ============================================================
add_heading_styled('4. Use Cases', 1)

add_heading_styled('Use Case 1: Sales Rep Self-Enablement', 2)
doc.add_paragraph(
    'A sales rep preparing for a customer meeting opens the Sales Bot and asks: '
    '"I need to sell cloud migration services to a financial services company competing against Accenture. '
    'What should I position and which slides should I use?" '
    'The bot responds with a tailored positioning strategy, recommends specific slides from the corpus, '
    'and can narrate the key points.'
)

add_heading_styled('Use Case 2: Live AI-Assisted Presentation', 2)
doc.add_paragraph(
    'In a Microsoft Teams meeting, the AI agent is added as a participant. A sales leader says '
    '"Present the One NTT Cloud Platform overview." The agent narrates the presentation with voice, '
    'advancing through slides. Audience members interrupt with questions and the agent answers in real-time.'
)

add_heading_styled('Use Case 3: Customer-Facing Product Demo', 2)
doc.add_paragraph(
    'A potential customer interacts directly with the Sales Bot through a web interface or Teams. '
    'They ask questions about NTT DATA\'s cloud services, pricing approach, and industry solutions. '
    'The bot provides accurate, consistent answers drawn from the latest product documentation.'
)

add_heading_styled('Use Case 4: New Hire Onboarding', 2)
doc.add_paragraph(
    'A new sales hire uses the bot to rapidly learn the product portfolio. They can ask open-ended '
    'questions, request presentations on specific topics, and get coaching on how to position products '
    'for different industries and competitive scenarios.'
)

doc.add_page_break()

# ============================================================
# 5. SOLUTION ARCHITECTURE
# ============================================================
add_heading_styled('5. Solution Architecture', 1)

doc.add_paragraph(
    'The recommended architecture combines Kore.ai\'s conversational AI platform with a RAG '
    '(Retrieval-Augmented Generation) pipeline to deliver accurate, up-to-date responses grounded '
    'in NTT DATA\'s actual product documentation.'
)

add_heading_styled('Architecture Components', 2)

components = [
    ('Conversational AI Layer (Kore.ai Platform)',
     'Natural language understanding, dialog management, voice synthesis, '
     'multi-channel deployment (web, Teams, voice). Kore.ai\'s enterprise agent platform '
     'provides the conversational interface and orchestration.'),
    ('Knowledge Ingestion Pipeline',
     'Automated ingestion from Confluence, Seismic, PowerPoint decks, Word documents, '
     'and internal portals. Documents are chunked, embedded, and indexed for retrieval.'),
    ('RAG Engine (Retrieval-Augmented Generation)',
     'Vector database storing document embeddings. On each query, relevant chunks are '
     'retrieved and passed to an LLM for grounded, accurate response generation.'),
    ('LLM Layer',
     'Large language model (configurable — Azure OpenAI, AWS Bedrock, or Google Vertex AI) '
     'for response generation, summarization, and presentation narration.'),
    ('Voice & Presentation Engine',
     'Text-to-speech with accent/language selection. Slide advancement and narration sync. '
     'Real-time Q&A interruption handling.'),
    ('Analytics & Reporting',
     'Usage tracking, CSAT surveys, sentiment analysis, meeting minutes generation, '
     'and presentation catalogue management.'),
]

for name, desc in components:
    add_bullet(desc, f'{name}: ')

add_heading_styled('Architecture Diagram (Conceptual)', 2)

doc.add_paragraph(
    '┌─────────────────────────────────────────────────────────────┐\n'
    '│  USERS: Sales Reps | Sales Engineers | Customers           │\n'
    '│  CHANNELS: Web Chat | MS Teams | Voice                     │\n'
    '└─────────────────────┬───────────────────────────────────────┘\n'
    '                      │\n'
    '                      ▼\n'
    '┌─────────────────────────────────────────────────────────────┐\n'
    '│  KORE.AI CONVERSATIONAL PLATFORM                           │\n'
    '│  • NLU / Dialog Management / Voice TTS                     │\n'
    '│  • Multi-channel orchestration                             │\n'
    '│  • Session management & analytics                          │\n'
    '└─────────────────────┬───────────────────────────────────────┘\n'
    '                      │\n'
    '                      ▼\n'
    '┌─────────────────────────────────────────────────────────────┐\n'
    '│  RAG ENGINE                                                │\n'
    '│  • Vector DB (embeddings of all product docs)              │\n'
    '│  • Semantic search + chunk retrieval                       │\n'
    '│  • LLM response generation (Azure/AWS/GCP)                │\n'
    '└─────────────────────┬───────────────────────────────────────┘\n'
    '                      │\n'
    '                      ▼\n'
    '┌─────────────────────────────────────────────────────────────┐\n'
    '│  KNOWLEDGE SOURCES                                         │\n'
    '│  • Confluence (auto-sync)                                  │\n'
    '│  • Seismic (slide library)                                 │\n'
    '│  • PowerPoint decks & Word docs                            │\n'
    '│  • Internal portals                                        │\n'
    '└─────────────────────────────────────────────────────────────┘\n'
).style = doc.styles['Normal']

doc.add_page_break()

# ============================================================
# 6. DELIVERY APPROACH
# ============================================================
add_heading_styled('6. Delivery Approach — RAG vs. Embedded Prompt', 1)

doc.add_paragraph(
    'A key architectural decision is how the bot accesses its knowledge. '
    'Below is a comparison of the three viable approaches.'
)

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['Approach', 'Pros', 'Cons', 'Best For']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

approaches = [
    ('RAG (Retrieval-Augmented Generation)',
     '• Always up-to-date\n• Scales to large doc corpus\n• Grounded answers with citations\n• Auto-sync with Confluence',
     '• More infrastructure\n• Needs vector DB\n• Slightly higher latency\n• Requires chunking tuning',
     'Production system with large, changing document corpus'),
    ('Giant Prompt with Embedded Data',
     '• Simple to build\n• Fast responses\n• No infrastructure needed\n• Great for POC/demo',
     '• Context window limits\n• Manual updates needed\n• Doesn\'t scale to large corpus\n• No auto-refresh',
     'Quick POC, small static knowledge base'),
    ('Hybrid (Embedded core + RAG for depth)',
     '• Core info always available\n• Deep search when needed\n• Good balance of speed/depth\n• Graceful degradation',
     '• Most complex to build\n• Two systems to maintain',
     'Best of both worlds for production'),
]

for row_idx, (approach, pros, cons, best) in enumerate(approaches, 1):
    table.rows[row_idx].cells[0].text = approach
    table.rows[row_idx].cells[1].text = pros
    table.rows[row_idx].cells[2].text = cons
    table.rows[row_idx].cells[3].text = best

doc.add_paragraph()

add_heading_styled('Recommendation', 2)
doc.add_paragraph(
    'For the POC (Phase 1): Use the Giant Prompt approach with embedded NTT DATA Cloud Services data. '
    'This allows rapid demonstration of the concept within days.\n\n'
    'For Production (Phase 2+): Migrate to RAG with Confluence auto-sync via the Kore.ai platform. '
    'This enables scaling to the full document corpus, automatic refresh, and grounded citations.'
)

doc.add_page_break()

# ============================================================
# 7. TECHNICAL SPECIFICATIONS
# ============================================================
add_heading_styled('7. Technical Specifications', 1)

add_heading_styled('Platform: Kore.ai XO Platform', 2)
add_bullet('Enterprise-grade conversational AI platform')
add_bullet('Pre-built NLU engine with multi-language support')
add_bullet('Voice gateway with TTS/STT (text-to-speech / speech-to-text)')
add_bullet('Microsoft Teams integration via Bot Framework')
add_bullet('RAG pipeline with vector search and LLM orchestration')
add_bullet('Analytics dashboard with sentiment analysis')
add_bullet('Enterprise security: SOC 2, HIPAA, GDPR compliant')

add_heading_styled('LLM Options', 2)
add_bullet('Azure OpenAI (GPT-4 / GPT-4o) — recommended for Microsoft ecosystem alignment')
add_bullet('AWS Bedrock (Claude) — alternative for AWS-aligned deployments')
add_bullet('Google Vertex AI (Gemini) — alternative for GCP-aligned deployments')
add_bullet('NTT DATA tsuzumi — for Japanese language support (0.6-7B parameters)')

add_heading_styled('Voice Capabilities', 2)
add_bullet('Real-time text-to-speech with accent selection')
add_bullet('Voice-to-voice LLM processing for natural conversation')
add_bullet('Multiple language support with translation')
add_bullet('Interruption handling for Q&A during presentations')

add_heading_styled('Integration Requirements', 2)
add_bullet('Microsoft Teams (Bot Framework SDK)')
add_bullet('Confluence API (content sync)')
add_bullet('Seismic API or file export (slide content)')
add_bullet('SSO / Azure AD for authentication')
add_bullet('Email / calendar for meeting minutes distribution')

doc.add_page_break()

# ============================================================
# 8. DATA REQUIREMENTS
# ============================================================
add_heading_styled('8. Data Requirements', 1)

doc.add_paragraph(
    'To build an effective POC and subsequent production system, the following data is needed from NTT DATA:'
)

add_heading_styled('POC Data (Minimum)', 2)
add_bullet('20-30 PowerPoint decks covering key cloud services products')
add_bullet('Product brochures and one-pagers for the Cloud Services Group')
add_bullet('Any existing presentation scripts or speaker notes')
add_bullet('Access to relevant Confluence spaces (read-only)')

add_heading_styled('Production Data (Full)', 2)
add_bullet('Complete Confluence space access for Cloud Services documentation')
add_bullet('Seismic content library access')
add_bullet('All product slide decks with notes sections')
add_bullet('Competitive positioning documents')
add_bullet('Pricing frameworks and proposal templates')
add_bullet('Case studies and customer success stories')
add_bullet('Technical architecture documents')

add_heading_styled('Data Processing', 2)
doc.add_paragraph(
    'Documents will be processed through the following pipeline:\n\n'
    '1. Ingestion — Extract text from PPTX, DOCX, PDF, Confluence pages\n'
    '2. Chunking — Split into semantic chunks (500-1000 tokens) with overlap\n'
    '3. Embedding — Generate vector embeddings using a text embedding model\n'
    '4. Indexing — Store in vector database with metadata (source, date, product area)\n'
    '5. Refresh — Automated sync on configurable schedule (daily/weekly)'
)

doc.add_page_break()

# ============================================================
# 9. INTEGRATION POINTS
# ============================================================
add_heading_styled('9. Integration Points', 1)

integrations = [
    ('Microsoft Teams', 'Primary delivery channel for live presentations and Q&A. Bot added as Teams member.', 'Phase 2'),
    ('Web Chat', 'Standalone web interface for self-service access. HTML/JS chatbot.', 'POC (Phase 1)'),
    ('Confluence', 'Knowledge source. Auto-sync new/modified/deleted content.', 'Phase 2'),
    ('Seismic', 'Slide library access for presentation content.', 'Phase 2'),
    ('Azure AD / SSO', 'Authentication for NTT DATA employees.', 'Phase 2'),
    ('Email / Calendar', 'Meeting minutes distribution, CSAT survey delivery.', 'Phase 3'),
    ('Analytics Dashboard', 'Usage reporting, sentiment tracking, CSAT scores.', 'Phase 2'),
]

table = doc.add_table(rows=len(integrations)+1, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['System', 'Purpose', 'Phase']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (system, purpose, phase) in enumerate(integrations, 1):
    table.rows[row_idx].cells[0].text = system
    table.rows[row_idx].cells[1].text = purpose
    table.rows[row_idx].cells[2].text = phase

doc.add_page_break()

# ============================================================
# 10. MVP / POC SCOPE
# ============================================================
add_heading_styled('10. MVP / POC Scope', 1)

doc.add_paragraph(
    'The POC demonstrates core feasibility with minimal infrastructure. '
    'Goal: Show Shreeshan and Wiley that the concept works and is worth investing in.'
)

add_heading_styled('POC Deliverables', 2)
add_numbered('HTML-based chatbot with embedded NTT DATA Cloud Services knowledge')
add_numbered('Natural language Q&A about cloud products, services, and capabilities')
add_numbered('Quick-action buttons for common topics')
add_numbered('Professional NTT DATA + Kore.ai branded interface')
add_numbered('Demonstration of voice capability (if time permits — browser TTS)')

add_heading_styled('POC Timeline', 2)
doc.add_paragraph(
    'With sample product documentation provided by NTT DATA:\n'
    '• Day 1-2: Ingest and index documentation via RAG pipeline\n'
    '• Day 3: Connect to Kore.ai platform with conversational flows\n'
    '• Day 4-5: Testing, tuning, and demo preparation\n\n'
    'Total: ~1 week from data receipt to working demo'
)

add_heading_styled('POC Success Criteria', 2)
add_bullet('Bot accurately answers questions about NTT DATA cloud services')
add_bullet('Responses are grounded in actual product documentation')
add_bullet('Interactive Q&A feels natural and conversational')
add_bullet('Stakeholders (Shreeshan, Wiley) confirm "yes, this is what we want"')

add_heading_styled('What the POC Does NOT Include', 2)
add_bullet('Microsoft Teams integration (Phase 2)')
add_bullet('Live voice presentation narration (Phase 2)')
add_bullet('Confluence auto-sync (Phase 2)')
add_bullet('CSAT surveys or sentiment analysis (Phase 3)')
add_bullet('Meeting minutes generation (Phase 3)')

doc.add_page_break()

# ============================================================
# 11. FUTURE ROADMAP
# ============================================================
add_heading_styled('11. Future Roadmap', 1)

phases = [
    ('Phase 1 — POC (Weeks 1-2)', [
        'Web-based chatbot with embedded knowledge',
        'NTT DATA Cloud Services knowledge base',
        'Natural language Q&A',
        'Branded interface (NTT DATA + Kore.ai)',
        'Demo to Shreeshan / Wiley for validation',
    ]),
    ('Phase 2 — MVP (Weeks 3-8)', [
        'Migrate to Kore.ai XO Platform',
        'RAG pipeline with vector database',
        'Confluence auto-sync integration',
        'Microsoft Teams bot deployment',
        'Voice TTS with accent/language selection',
        'Basic analytics and usage tracking',
        'SSO / Azure AD authentication',
    ]),
    ('Phase 3 — Production (Weeks 9-16)', [
        'Presentation catalogue with version control',
        'Live presentation mode (slide narration)',
        'Meeting minutes auto-generation',
        'CSAT survey integration',
        'Sentiment analysis on conversations',
        'Usage reporting dashboard',
        'Multi-language support',
        'Seismic integration',
    ]),
    ('Phase 4 — Scale (Ongoing)', [
        'Expand to additional product lines beyond Cloud Services',
        'Customer-facing deployment option',
        'Advanced competitive positioning agent',
        'Proposal generation assistance',
        'Integration with CRM (Salesforce)',
    ]),
]

for phase_name, items in phases:
    add_heading_styled(phase_name, 2)
    for item in items:
        add_bullet(item)

doc.add_page_break()

# ============================================================
# 12. OPEN QUESTIONS
# ============================================================
add_heading_styled('12. Open Questions', 1)

questions = [
    ('Volume & Usage', 'How many sessions per month are expected? Initial estimate: 200-500 users, but frequency per user TBD.'),
    ('Pricing Model', 'Is this a per-session cost, platform license, or managed service? Needs alignment with the $3.5M budget framework.'),
    ('Build vs. Buy', 'Does NTT DATA want to build on Kore.ai platform themselves, or have Kore.ai build and deliver?'),
    ('Data Access', 'When can NTT DATA provide 20-30 sample PowerPoint/product docs for the POC?'),
    ('Audience', 'Is this for internal sales enablement only, or also customer-facing? Discovery suggests both.'),
    ('Voice Priority', 'How critical is voice narration for the POC vs. production? Can POC be text-only?'),
    ('Competitive Landscape', 'Are other vendors being evaluated for this solution?'),
    ('Success Metrics', 'How will NTT DATA measure ROI? Time saved per rep? Presentation quality? Usage rates?'),
    ('Security Requirements', 'What are the data handling and security requirements for product documentation?'),
    ('Confluence Access', 'Which Confluence spaces are in scope? Who grants access?'),
]

table = doc.add_table(rows=len(questions)+1, cols=2)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Topic', 'Question / Notes']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, (topic, question) in enumerate(questions, 1):
    table.rows[row_idx].cells[0].text = topic
    table.rows[row_idx].cells[1].text = question

doc.add_page_break()

# ============================================================
# 13. APPENDIX
# ============================================================
add_heading_styled('13. Appendix — NTT DATA Cloud Services Knowledge Base', 1)

doc.add_paragraph(
    'The following data was crawled from public NTT DATA web properties on 2026-04-01 '
    'and used to seed the POC chatbot. This represents the type of content that would be '
    'ingested from internal documentation in the production system.'
)

add_heading_styled('Company Overview', 2)
doc.add_paragraph(
    'NTT DATA is one of the world\'s largest IT services providers, offering full-stack cloud services '
    'with a platform-first approach powered by AI. They serve 700+ global clients and have completed '
    '600+ public-cloud IT transformation projects. Key value proposition: 30% faster to market, '
    '30% less cost.'
)

add_heading_styled('Cloud Services Portfolio', 2)
add_bullet('Cloud Strategy and Transformation')
add_bullet('Cloud Architecture and Modernization')
add_bullet('Cloud Platforms (Industry, Sovereign, Public, Private, Containers)')
add_bullet('Cloud Optimization')
add_bullet('Cloud Management (Run Ops, Agile PODs, SRE, CCoE, FinOps)')
add_bullet('Cloud Implementation (Migration, App Transformation, Landing Zones)')

add_heading_styled('AI & GenAI Platforms', 2)
add_bullet('NTT DATA Generative AI Technology Hub (TechHub) — 40% faster deployment')
add_bullet('tsuzumi Language Model — 0.6-7B params, English/Japanese')
add_bullet('Smart AI Agent Ecosystem — healthcare, automotive, finance, supply chain agents')
add_bullet('Workplace Smart AI Agent Suite — Knowledge.ai, Quality.ai, Copilot')

add_heading_styled('Nucleus Intelligent Enterprise Platform', 2)
add_bullet('Command Center — unifies ~100 reports into single interface')
add_bullet('Runbook Automation — multi-step, multi-vendor lifecycle management')
add_bullet('Intelligent Patch Management — automated cross-platform updates')
add_bullet('Configuration Manager — role-based infrastructure provisioning')

add_heading_styled('Key Partnerships', 2)
add_bullet('Microsoft: Global GSI, Azure Expert MSP')
add_bullet('AWS: Premier Consulting Partner, Strategic Collaboration (Jan 2026)')
add_bullet('Google Cloud: Premier Partner')
add_bullet('Kore.ai: Enterprise agent platform')
add_bullet('OpenAI: Center of Excellence')

add_heading_styled('Sources', 2)
sources = [
    'https://services.global.ntt/en-us/services-and-products/cloud',
    'https://us.nttdata.com/en/services/cloud-and-it-infrastructure-services',
    'https://us.nttdata.com/en/services/data-and-artificial-intelligence',
    'https://us.nttdata.com/en/services/generative-ai/genai-platforms',
    'https://us.nttdata.com/en/services/cloud-and-it-infrastructure-services/cloud-management',
    'https://us.nttdata.com/en/services/nucleus-intelligent-enterprise-platform/cloud-and-hybrid-infrastructure',
    'https://services.global.ntt/en-us/newsroom/ntt-data-unveils-smart-ai-agent-ecosystem',
    'https://us.nttdata.com/en/services/cybersecurity-services/cloud-security',
]
for src in sources:
    add_bullet(src)

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/michaelashe/Documents/NTT_Data_SalesBot/OneNTTData_SalesBot_RequirementsAndDesign_v01_April2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
